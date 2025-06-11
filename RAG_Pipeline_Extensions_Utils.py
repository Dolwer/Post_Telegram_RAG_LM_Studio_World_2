import os
import json
import csv
import requests
from typing import List, Dict, Optional, Generator
from pathlib import Path
import asyncio
import aiohttp
from dataclasses import dataclass
import xml.etree.ElementTree as ET
from urllib.parse import urlparse, urljoin
import time
import hashlib

# Попытка импорта дополнительных библиотек (установить при необходимости)
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("PyPDF2 не установлен. Функциональность PDF недоступна.")

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    print("BeautifulSoup4 не установлен. Функциональность парсинга HTML недоступна.")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx не установлен. Функциональность DOCX недоступна.")

from advanced_rag_pipeline import Document, AdvancedRAGPipeline

class DataIngestionManager:
    """Менеджер для загрузки данных из различных источников"""
    
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
    
    def load_from_text_file(self, filepath: str, encoding: str = 'utf-8') -> str:
        """Загрузка текста из файла"""
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Ошибка чтения файла {filepath}: {e}")
    
    def load_from_pdf(self, filepath: str) -> str:
        """Загрузка текста из PDF файла"""
        if not PDF_AVAILABLE:
            raise Exception("PyPDF2 не установлен. Установите: pip install PyPDF2")
        
        try:
            text = ""
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Ошибка чтения PDF {filepath}: {e}")
    
    def load_from_docx(self, filepath: str) -> str:
        """Загрузка текста из DOCX файла"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx не установлен. Установите: pip install python-docx")
        
        try:
            doc = docx.Document(filepath)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Ошибка чтения DOCX {filepath}: {e}")
    
    def load_from_csv(self, filepath: str, text_columns: List[str]) -> List[Dict]:
        """Загрузка данных из CSV файла"""
        try:
            documents = []
            with open(filepath, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for i, row in enumerate(reader):
                    # Объединяем указанные колонки в один текст
                    content_parts = []
                    for col in text_columns:
                        if col in row and row[col]:
                            content_parts.append(str(row[col]))
                    
                    if content_parts:
                        content = " ".join(content_parts)
                        # Остальные колонки как метаданные
                        metadata = {k: v for k, v in row.items() if k not in text_columns}
                        
                        documents.append({
                            'id': f"csv_row_{i}",
                            'content': content,
                            'metadata': metadata
                        })
            
            return documents
        except Exception as e:
            raise Exception(f"Ошибка чтения CSV {filepath}: {e}")
    
    def load_from_json(self, filepath: str, content_field: str, id_field: Optional[str] = None) -> List[Dict]:
        """Загрузка данных из JSON файла"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            documents = []
            if isinstance(data, list):
                for i, item in enumerate(data):
                    if content_field in item:
                        doc_id = item.get(id_field, f"json_item_{i}") if id_field else f"json_item_{i}"
                        content = str(item[content_field])
                        metadata = {k: v for k, v in item.items() if k not in [content_field, id_field]}
                        
                        documents.append({
                            'id': doc_id,
                            'content': content,
                            'metadata': metadata
                        })
            
            return documents
        except Exception as e:
            raise Exception(f"Ошибка чтения JSON {filepath}: {e}")
    
    def load_from_url(self, url: str, timeout: int = 30) -> str:
        """Загрузка контента с URL"""
        try:
            response = requests.get(url, timeout=timeout)
            response.raise_for_status()
            
            # Если это HTML, пытаемся извлечь текст
            if 'text/html' in response.headers.get('content-type', ''):
                if BS4_AVAILABLE:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    # Удаляем script и style теги
                    for script in soup(["script", "style"]):
                        script.decompose()
                    return soup.get_text()
                else:
                    return response.text
            else:
                return response.text
                
        except Exception as e:
            raise Exception(f"Ошибка загрузки URL {url}: {e}")
    
    async def load_from_urls_async(self, urls: List[str], timeout: int = 30) -> List[Dict]:
        """Асинхронная загрузка множества URL"""
        async def fetch_url(session, url):
            try:
                async with session.get(url, timeout=timeout) as response:
                    content = await response.text()
                    if 'text/html' in response.headers.get('content-type', ''):
                        if BS4_AVAILABLE:
                            soup = BeautifulSoup(content, 'html.parser')
                            for script in soup(["script", "style"]):
                                script.decompose()
                            content = soup.get_text()
                    
                    return {
                        'id': hashlib.md5(url.encode()).hexdigest(),
                        'content': content,
                        'metadata': {'source_url': url, 'status': 'success'}
                    }
            except Exception as e:
                return {
                    'id': hashlib.md5(url.encode()).hexdigest(),
                    'content': "",
                    'metadata': {'source_url': url, 'status': 'error', 'error': str(e)}
                }
        
        async with aiohttp.ClientSession() as session:
            tasks = [fetch_url(session, url) for url in urls]
            results = await asyncio.gather(*tasks)
            return [r for r in results if r['content']]  # Возвращаем только успешные

class RAGAnalytics:
    """Аналитика и мониторинг RAG системы"""
    
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.query_log = []
    
    def log_query(self, query: str, results_count: int, processing_time: float):
        """Логирование запроса"""
        self.query_log.append({
            'timestamp': time.time(),
            'query': query,
            'results_count': results_count,
            'processing_time': processing_time
        })
    
    def get_query_stats(self) -> Dict:
        """Статистика по запросам"""
        if not self.query_log:
            return {"message": "Нет данных по запросам"}
        
        processing_times = [log['processing_time'] for log in self.query_log]
        results_counts = [log['results_count'] for log in self.query_log]
        
        return {
            'total_queries': len(self.query_log),
            'avg_processing_time': sum(processing_times) / len(processing_times),
            'max_processing_time': max(processing_times),
            'min_processing_time': min(processing_times),
            'avg_results_count': sum(results_counts) / len(results_counts),
            'queries_per_hour': len([q for q in self.query_log if time.time() - q['timestamp'] < 3600])
        }
    
    def analyze_collection_content(self) -> Dict:
        """Анализ контента коллекции"""
        try:
            # Получаем все документы
            all_data = self.rag.collection.get()
            documents = all_data.get('documents', [])
            metadatas = all_data.get('metadatas', [])
            
            if not documents:
                return {"message": "Коллекция пуста"}
            
            # Базовая статистика
            total_docs = len(documents)
            total_chars = sum(len(doc) for doc in documents)
            avg_doc_length = total_chars / total_docs
            
            # Анализ метаданных
            categories = {}
            languages = {}
            
            for metadata in metadatas:
                if metadata:
                    if 'category' in metadata:
                        cat = metadata['category']
                        categories[cat] = categories.get(cat, 0) + 1
                    
                    if 'language' in metadata:
                        lang = metadata['language']
                        languages[lang] = languages.get(lang, 0) + 1
            
            return {
                'total_documents': total_docs,
                'total_characters': total_chars,
                'average_document_length': avg_doc_length,
                'categories_distribution': categories,
                'languages_distribution': languages,
                'longest_document': max(len(doc) for doc in documents),
                'shortest_document': min(len(doc) for doc in documents)
            }
            
        except Exception as e:
            return {"error": f"Ошибка анализа: {e}"}

class RAGWebInterface:
    """Простой веб-интерфейс для RAG системы"""
    
    def __init__(self, rag_pipeline: AdvancedRAGPipeline, analytics: RAGAnalytics):
        self.rag = rag_pipeline
        self.analytics = analytics
    
    def generate_html_interface(self) -> str:
        """Генерация HTML интерфейса"""
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>RAG Pipeline Interface</title>
            <meta charset="UTF-8">
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; background-color: #f5f5f5; }
                .container { max-width: 1200px; margin: 0 auto; background: white; padding: 20px; border-radius: 8px; }
                .search-box { width: 70%; padding: 10px; font-size: 16px; border: 1px solid #ddd; border-radius: 4px; }
                .search-btn { padding: 10px 20px; background: #007bff; color: white; border: none; border-radius: 4px; cursor: pointer; }
                .result { margin: 15px 0; padding: 15px; border: 1px solid #ddd; border-radius: 4px; background: #f9f9f9; }
                .score { color: #666; font-size: 14px; }
                .metadata { color: #888; font-size: 12px; margin-top: 5px; }
                .stats { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin: 20px 0; }
                .stat-card { padding: 15px; background: #e9ecef; border-radius: 4px; text-align: center; }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>🔍 RAG Pipeline Interface</h1>
                
                <div class="search-section">
                    <input type="text" class="search-box" id="searchInput" placeholder="Введите ваш запрос...">
                    <button class="search-btn" onclick="performSearch()">Поиск</button>
                </div>
                
                <div id="results"></div>
                
                <h2>📊 Статистика системы</h2>
                <div class="stats" id="stats">
                    <div class="stat-card">
                        <h3>Запросы</h3>
                        <div id="queryCount">Загрузка...</div>
                    </div>
                    <div class="stat-card">
                        <h3>Среднее время</h3>
                        <div id="avgTime">Загрузка...</div>
                    </div>
                </div>
                
                <h2>📈 Управление данными</h2>
                <div style="margin: 20px 0;">
                    <button onclick="exportData()" style="margin-right: 10px; padding: 8px 16px; background: #28a745; color: white; border: none; border-radius: 4px;">Экспорт данных</button>
                    <button onclick="clearCollection()" style="padding: 8px 16px; background: #dc3545; color: white; border: none; border-radius: 4px;">Очистить коллекцию</button>
                </div>
            </div>
            
            <script>
                async function performSearch() {
                    const query = document.getElementById('searchInput').value;
                    if (!query.trim()) return;
                    
                    const resultsDiv = document.getElementById('results');
                    resultsDiv.innerHTML = '<div>Поиск...</div>';
                    
                    try {
                        // Здесь будет API вызов к серверу RAG
                        // Для демонстрации показываем статичные результаты
                        setTimeout(() => {
                            resultsDiv.innerHTML = `
                                <h3>Результаты поиска для: "${query}"</h3>
                                <div class="result">
                                    <strong>Документ 1</strong>
                                    <div class="score">Релевантность: 0.85</div>
                                    <p>Пример найденного контента...</p>
                                    <div class="metadata">Метаданные: category=ai, timestamp=2024-01-01</div>
                                </div>
                            `;
                        }, 1000);
                    } catch (error) {
                        resultsDiv.innerHTML = '<div style="color: red;">Ошибка поиска: ' + error.message + '</div>';
                    }
                }
                
                function loadStats() {
                    // Здесь будет загрузка реальной статистики
                    document.getElementById('docCount').textContent = '150';
                    document.getElementById('queryCount').textContent = '45';
                    document.getElementById('avgTime').textContent = '0.3s';
                }
                
                function exportData() {
                    alert('Функция экспорта будет реализована на сервере');
                }
                
                function clearCollection() {
                    if (confirm('Вы уверены, что хотите очистить коллекцию?')) {
                        alert('Функция очистки будет реализована на сервере');
                    }
                }
                
                // Загрузка статистики при загрузке страницы
                window.onload = loadStats;
                
                // Поиск по Enter
                document.getElementById('searchInput').addEventListener('keypress', function(e) {
                    if (e.key === 'Enter') {
                        performSearch();
                    }
                });
            </script>
        </body>
        </html>
        """
        return html_template
    
    def save_interface(self, filepath: str = "rag_interface.html"):
        """Сохранение HTML интерфейса в файл"""
        html_content = self.generate_html_interface()
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"HTML интерфейс сохранен в {filepath}")

class RAGBenchmarking:
    """Бенчмаркинг и тестирование RAG системы"""
    
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
    
    def create_test_dataset(self, size: int = 100) -> List[Document]:
        """Создание тестового датасета"""
        test_topics = [
            "машинное обучение", "искусственный интеллект", "нейронные сети",
            "векторные базы данных", "обработка естественного языка",
            "компьютерное зрение", "глубокое обучение", "анализ данных",
            "программирование на Python", "веб-разработка"
        ]
        
        documents = []
        for i in range(size):
            topic = test_topics[i % len(test_topics)]
            content = self._generate_synthetic_content(topic, i)
            
            doc = Document(
                id=f"test_doc_{i}",
                content=content,
                metadata={
                    "topic": topic,
                    "test_id": i,
                    "synthetic": True
                }
            )
            documents.append(doc)
        
        return documents
    
    def _generate_synthetic_content(self, topic: str, doc_id: int) -> str:
        """Генерация синтетического контента для тестов"""
        templates = [
            f"{topic} является важной областью современной технологии. Документ номер {doc_id} содержит подробную информацию о применении {topic} в различных сферах.",
            f"В этом документе рассматриваются основные принципы {topic}. Это руководство номер {doc_id} поможет понять ключевые концепции.",
            f"Практическое применение {topic} демонстрируется в примере {doc_id}. Здесь представлены лучшие практики и методы."
        ]
        
        base_content = templates[doc_id % len(templates)]
        # Добавляем дополнительный контент для разнообразия
        additional = f" Дополнительные детали включают технические аспекты, примеры использования и рекомендации экспертов в области {topic}."
        
        return base_content + additional
    
    def benchmark_search_performance(self, queries: List[str], n_runs: int = 10) -> Dict:
        """Бенчмарк производительности поиска"""
        results = {
            'queries': [],
            'avg_time': 0,
            'total_time': 0,
            'fastest_query': None,
            'slowest_query': None
        }
        
        total_time = 0
        fastest_time = float('inf')
        slowest_time = 0
        
        for query in queries:
            query_times = []
            
            for _ in range(n_runs):
                start_time = time.time()
                search_results = self.rag.search(query, n_results=5)
                end_time = time.time()
                
                query_time = end_time - start_time
                query_times.append(query_time)
                total_time += query_time
            
            avg_query_time = sum(query_times) / len(query_times)
            min_query_time = min(query_times)
            max_query_time = max(query_times)
            
            if min_query_time < fastest_time:
                fastest_time = min_query_time
                results['fastest_query'] = query
            
            if max_query_time > slowest_time:
                slowest_time = max_query_time
                results['slowest_query'] = query
            
            results['queries'].append({
                'query': query,
                'avg_time': avg_query_time,
                'min_time': min_query_time,
                'max_time': max_query_time,
                'runs': n_runs
            })
        
        results['avg_time'] = total_time / (len(queries) * n_runs)
        results['total_time'] = total_time
        results['fastest_time'] = fastest_time
        results['slowest_time'] = slowest_time
        
        return results
    
    def evaluate_retrieval_quality(self, test_queries: List[Dict]) -> Dict:
        """Оценка качества поиска
        test_queries: [{'query': str, 'relevant_doc_ids': List[str]}]
        """
        metrics = {
            'precision_at_k': [],
            'recall_at_k': [],
            'average_precision': []
        }
        
        for test_case in test_queries:
            query = test_case['query']
            relevant_ids = set(test_case['relevant_doc_ids'])
            
            # Получаем результаты поиска
            results = self.rag.search(query, n_results=10)
            retrieved_ids = [r.document_id for r in results]
            
            # Вычисляем метрики для разных K
            for k in [1, 3, 5, 10]:
                if k <= len(retrieved_ids):
                    retrieved_k = set(retrieved_ids[:k])
                    
                    # Precision@K
                    precision = len(retrieved_k & relevant_ids) / k
                    
                    # Recall@K
                    recall = len(retrieved_k & relevant_ids) / len(relevant_ids) if relevant_ids else 0
                    
                    metrics['precision_at_k'].append(precision)
                    metrics['recall_at_k'].append(recall)
            
            # Average Precision
            ap = self._calculate_average_precision(retrieved_ids, relevant_ids)
            metrics['average_precision'].append(ap)
        
        # Усредняем метрики
        final_metrics = {
            'mean_precision_at_k': sum(metrics['precision_at_k']) / len(metrics['precision_at_k']) if metrics['precision_at_k'] else 0,
            'mean_recall_at_k': sum(metrics['recall_at_k']) / len(metrics['recall_at_k']) if metrics['recall_at_k'] else 0,
            'mean_average_precision': sum(metrics['average_precision']) / len(metrics['average_precision']) if metrics['average_precision'] else 0
        }
        
        return final_metrics
    
    def _calculate_average_precision(self, retrieved_ids: List[str], relevant_ids: set) -> float:
        """Вычисление Average Precision"""
        if not relevant_ids:
            return 0.0
        
        precisions = []
        relevant_count = 0
        
        for i, doc_id in enumerate(retrieved_ids):
            if doc_id in relevant_ids:
                relevant_count += 1
                precision = relevant_count / (i + 1)
                precisions.append(precision)
        
        return sum(precisions) / len(relevant_ids) if precisions else 0.0

# Дополнительные утилиты
class RAGConfigManager:
    """Менеджер конфигурации RAG системы"""
    
    def __init__(self, config_path: str = "rag_config.json"):
        self.config_path = config_path
        self.default_config = {
            "embedding_model": "all-MiniLM-L6-v2",
            "chunk_size": 500,
            "chunk_overlap": 50,
            "collection_name": "default_rag",
            "persist_directory": "./rag_storage",
            "search_results_limit": 5,
            "similarity_threshold": 0.3,
            "batch_size": 100,
            "logging_level": "INFO"
        }
        self.config = self.load_config()
    
    def load_config(self) -> Dict:
        """Загрузка конфигурации"""
        try:
            if os.path.exists(self.config_path):
                with open(self.config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                # Дополняем недостающими параметрами
                for key, value in self.default_config.items():
                    if key not in config:
                        config[key] = value
                return config
            else:
                return self.default_config.copy()
        except Exception as e:
            print(f"Ошибка загрузки конфигурации: {e}")
            return self.default_config.copy()
    
    def save_config(self):
        """Сохранение конфигурации"""
        try:
            with open(self.config_path, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Ошибка сохранения конфигурации: {e}")
    
    def update_config(self, updates: Dict):
        """Обновление конфигурации"""
        self.config.update(updates)
        self.save_config()
    
    def get(self, key: str, default=None):
        """Получение значения конфигурации"""
        return self.config.get(key, default)

# Пример интеграции всех компонентов
class FullRAGSystem:
    """Полная RAG система со всеми компонентами"""
    
    def __init__(self, config_path: str = "rag_config.json"):
        # Загружаем конфигурацию
        self.config_manager = RAGConfigManager(config_path)
        config = self.config_manager.config
        
        # Инициализируем основную систему
        self.rag = AdvancedRAGPipeline(
            collection_name=config['collection_name'],
            persist_directory=config['persist_directory'],
            embedding_model=config['embedding_model']
        )
        
        # Инициализируем дополнительные компоненты
        self.data_manager = DataIngestionManager(self.rag)
        self.analytics = RAGAnalytics(self.rag)
        self.web_interface = RAGWebInterface(self.rag, self.analytics)
        self.benchmarking = RAGBenchmarking(self.rag)
        
        print("🚀 Полная RAG система инициализирована!")
    
    def ingest_directory(self, directory_path: str, file_extensions: List[str] = None):
        """Загрузка всех файлов из директории"""
        if file_extensions is None:
            file_extensions = ['.txt', '.pdf', '.docx', '.json', '.csv']
        
        directory = Path(directory_path)
        if not directory.exists():
            raise Exception(f"Директория {directory_path} не существует")
        
        documents = []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_extensions:
                try:
                    if file_path.suffix.lower() == '.txt':
                        content = self.data_manager.load_from_text_file(str(file_path))
                        doc = Document(
                            id=str(file_path.stem),
                            content=content,
                            metadata={'source_file': str(file_path), 'file_type': 'text'}
                        )
                        documents.append(doc)
                    
                    elif file_path.suffix.lower() == '.pdf' and PDF_AVAILABLE:
                        content = self.data_manager.load_from_pdf(str(file_path))
                        doc = Document(
                            id=str(file_path.stem),
                            content=content,
                            metadata={'source_file': str(file_path), 'file_type': 'pdf'}
                        )
                        documents.append(doc)
                    
                    elif file_path.suffix.lower() == '.docx' and DOCX_AVAILABLE:
                        content = self.data_manager.load_from_docx(str(file_path))
                        doc = Document(
                            id=str(file_path.stem),
                            content=content,
                            metadata={'source_file': str(file_path), 'file_type': 'docx'}
                        )
                        documents.append(doc)
                    
                    print(f"✅ Обработан файл: {file_path.name}")
                    
                except Exception as e:
                    print(f"❌ Ошибка обработки файла {file_path.name}: {e}")
        
        if documents:
            self.rag.add_documents_batch(documents)
            print(f"📚 Загружено {len(documents)} документов из {directory_path}")
        else:
            print("⚠️ Не найдено подходящих файлов для загрузки")
    
    def create_web_interface(self, output_path: str = "rag_interface.html"):
        """Создание веб-интерфейса"""
        self.web_interface.save_interface(output_path)
        print(f"🌐 Веб-интерфейс создан: {output_path}")
    
    def run_performance_test(self):
        """Запуск тестов производительности"""
        test_queries = [
            "машинное обучение",
            "векторные базы данных",
            "обработка текста",
            "искусственный интеллект",
            "Python программирование"
        ]
        
        print("🧪 Запуск тестов производительности...")
        results = self.benchmarking.benchmark_search_performance(test_queries)
        
        print(f"📊 Результаты тестирования:")
        print(f"   Среднее время поиска: {results['avg_time']:.3f}s")
        print(f"   Самый быстрый запрос: {results['fastest_query']} ({results['fastest_time']:.3f}s)")
        print(f"   Самый медленный запрос: {results['slowest_query']} ({results['slowest_time']:.3f}s)")
        
        return results
    
    def get_system_status(self) -> Dict:
        """Получение статуса системы"""
        collection_stats = self.rag.get_collection_stats()
        query_stats = self.analytics.get_query_stats()
        content_analysis = self.analytics.analyze_collection_content()
        
        return {
            'collection_stats': collection_stats,
            'query_stats': query_stats,
            'content_analysis': content_analysis,
            'config': self.config_manager.config
        }

# Демонстрация использования всех компонентов
if __name__ == "__main__":
    print("🔧 Инициализация полной RAG системы...")
    
    # Создаем полную систему
    full_system = FullRAGSystem()
    
    # Создаем тестовые данные
    test_docs = [
        Document(
            id="advanced_doc_1",
            content="Расширенные возможности RAG включают обработку различных форматов файлов, асинхронную загрузку данных и аналитику производительности.",
            metadata={"category": "advanced", "topic": "rag_features"}
        ),
        Document(
            id="advanced_doc_2", 
            content="Система мониторинга позволяет отслеживать производительность поиска, анализировать запросы пользователей и оптимизировать работу векторной базы данных.",
            metadata={"category": "monitoring", "topic": "analytics"}
        )
    ]
    
    # Добавляем документы
    full_system.rag.add_documents_batch(test_docs)
    
    # Тестируем поиск
    print("\n🔍 Тестирование поиска...")
    results = full_system.rag.search("аналитика производительности RAG", n_results=2)
    for i, result in enumerate(results):
        print(f"Результат {i+1}: {result.content[:100]}... (схожесть: {result.similarity_score:.3f})")
    
    # Получаем статус системы
    print("\n📊 Статус системы:")
    status = full_system.get_system_status()
    print(json.dumps(status, indent=2, ensure_ascii=False))
    
    # Создаем веб-интерфейс
    full_system.create_web_interface()
    
    # Запускаем тесты производительности
    performance_results = full_system.run_performance_test()
    
    print("\n✅ Демонстрация завершена! Система готова к использованию.")card">
                        <h3>Документы</h3>
                        <div id="docCount">Загрузка...</div>
                    </div>
                    <div class="stat-
